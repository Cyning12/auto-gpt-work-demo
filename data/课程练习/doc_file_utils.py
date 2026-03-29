"""
Word 文档（.doc / .docx）解析与导出 Markdown；供迪士尼 RAG 等练习复用。

**磁盘上的原始形态**

- **.docx**：本质是 **ZIP**，内含 ``word/document.xml`` 等 **OOXML**（XML 文件树），不是单一扁平 XML 字符串。
- **.doc**（Word 97–2003）：**OLE 复合文档**（二进制容器），**不是** XML；本模块通过外部命令抽取纯文本。

**解析后在 Python 里长什么样**

统一为 ``list[dict[str, str]]``：每项 ``{"type": "text"|"table", "content": str}``。``.docx`` 常为多块（段落 + 可选 Markdown 表）；
``.doc`` 多为 **一个** ``text`` 大块（全文），见 ``parse_docx`` / ``parse_doc`` 文档。

**``doc.element`` 是不是 XML**

``python-docx`` 的 ``Document.element`` 是 **OXML 元素对象**（如 ``CT_Document``，继承 ``BaseOxmlElement``），底层基于 **lxml**，
可用 ``tag``、``findall``、``iterchildren`` 等，与「整份文档的 XML 文本」不是同一概念。

**``.doc`` 后端：textutil / antiword 与服务器**

- **textutil**：**仅 macOS** 自带，Linux/Windows 服务器上通常 **没有**。
- **antiword**：**不是** Mac 专用；许多 **Linux** 镜像可通过包管理安装（例如 Debian/Ubuntu：``apt install antiword``），
  适合作为无 macOS 时的首选；macOS 也可用 Homebrew 安装。
- 若生产环境两者都不可用：可改用 **LibreOffice 无头**（``soffice --headless --convert-to txt doc``）、``catdoc`` 等；
  或先把 ``.doc`` **转成** ``.docx`` 再调用 ``parse_docx``（结构信息更完整）。

**RAG 切片**

- ``split_parsed_word_blocks``：对已解析块列表做 ``RecursiveCharacterTextSplitter`` 细切。
- ``build_chunks_for_word_document``：单文件解析 + 切片，输出与 ``pdf_file_utils.build_chunks_for_pdf`` 同形列表（另含 ``block_index`` / ``block_type``）。
- ``markdown_from_word_export_to_parsed_blocks``：将 ``chunks_to_markdown`` 写出的 MD **还原**为 ``list[{"type","content"}]``，再供切片。
- ``build_chunks_from_parsed_markdown_file`` / ``export_parsed_markdown_chunks_for_classified_docs``：对已导出 MD 批量切片并写 JSON。
"""

from __future__ import annotations

import hashlib
import json
import re
import subprocess
from pathlib import Path
from typing import Any

from langchain_text_splitters import RecursiveCharacterTextSplitter

# 与 pdf_file_utils / 课程向量库默认元数据对齐
DEFAULT_WORD_CHUNK_METADATA_DEPARTMENT = "company"
DEFAULT_WORD_CHUNK_METADATA_UPDATE_TIME = "2026-03-27"


def parse_docx(file_path: str | Path) -> list[dict[str, str]]:
    """
    解析 DOCX（OOXML），按 **文档主体中元素出现顺序** 抽出段落与表格，供后续拼 Markdown 或 RAG 分块。

    不依赖 ``python-docx`` 的 ``doc.paragraphs`` / ``doc.tables`` 与 body 顺序一一对应关系，
    而是直接遍历 ``document`` 根下 ``body`` 的子节点，避免复杂版式下表格与段落错位。

    **返回值**

    ``list[dict[str, str]]``，每项为一块内容（与 ``parse_doc`` 的键相同，便于 ``parse_word_document`` / ``chunks_to_markdown`` 统一处理；但 ``.docx`` 可多块且含表，见 ``parse_doc`` 说明）：

    - ``{"type": "text", "content": "..."}``：一段连续正文（已 ``strip``，空段不落库）。
    - ``{"type": "table", "content": "..."}``：整张表转成的 **GitHub 风格 Markdown 表格**（首行作表头，其余作数据行）。

    **格式样例**（节选）::

        [
            {"type": "text", "content": "第一章 总则"},
            {"type": "text", "content": "第一条 本规定适用于……"},
            {"type": "table", "content": "| 项目 | 说明 |\\n| --- | --- |\\n| A | 详情 |"},
        ]

    **函数内各段代码作用（与实现顺序一致）**

    1. **路径与后缀**：规范为 ``Path``，非 ``.docx`` 直接报错，避免误用 ``.doc`` 走 XML 分支。
    2. **延迟导入 ``python-docx``**：仅在调用时导入；若环境中误装 PyPI 上的旧包 ``docx`` 会与 ``python-docx`` 冲突，此处抛出明确 ``ImportError`` 提示卸载/重装。
    3. **打开文档与容器**：``Document`` 读入文件；``doc.element`` 为文档根 **OXML 元素**（非 XML 字符串）；
       ``doc.element.body`` 即 ``w:body`` 子树；``content_chunks`` 按阅读顺序累积块；``ns`` 为 WordprocessingML 命名空间前缀，
       供子元素上的 ``findall`` 使用 ``w:p``、``w:tbl`` 等。
    4. **按 body 子节点顺序遍历**：对每个子元素看 XML ``tag``：
       - **段落 ``w:p``**：在段落子树内收集所有 ``w:t`` 文本节点并拼接，得到该段完整字符串（含行内格式合并后的纯文本，不保留粗体等结构）。
       - **表格 ``w:tbl``**：取出所有行 ``w:tr``；用嵌套函数按行抽单元格 ``w:tc`` 内文本，第一行作 Markdown 表头与分隔行，后续行作数据行；非空则追加 ``type=table`` 块。
    5. **返回**：顺序列表，与原文档中块出现顺序一致，便于 ``chunks_to_markdown`` 按序输出。

    :param file_path: ``.docx`` 文件路径。
    :raises ValueError: 后缀不是 ``.docx``。
    :raises ImportError: 缺少或错误的 ``python-docx`` 依赖。
    """
    # 1) 路径与后缀校验
    p = Path(file_path)
    if p.suffix.lower() != ".docx":
        raise ValueError(f"仅支持解析 .docx 文件，当前为: {p.name}")
    # 2) 依赖：须为 python-docx（包名 python-docx，import 名 docx）
    try:
        from docx import Document as DocxDocument
    except Exception as e:
        raise ImportError(
            "DOCX 解析依赖异常：请卸载旧版 docx 并安装 python-docx（pip uninstall docx && pip install python-docx）"
        ) from e

    # 3) 打开文档、顺序块列表、WordprocessingML 命名空间（w →  URI）
    doc = DocxDocument(str(p))
    content_chunks: list[dict[str, str]] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # 4) 严格按 body 下子节点顺序处理段落 / 表格
    for child in doc.element.body.iterchildren():
        tag = getattr(child, "tag", None)
        if tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
            # 段落：汇总本段内所有 w:t 文本
            texts = [t.text for t in child.findall(".//w:t", ns) if t.text]
            paragraph_text = "".join(texts).strip()
            if paragraph_text:
                content_chunks.append({"type": "text", "content": paragraph_text})
        elif tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl":
            # 表格：按行解析单元格，转 Markdown 管道表
            rows = child.findall("./w:tr", ns)
            if not rows:
                continue

            md_table: list[str] = []

            def _row_cells(tr: Any) -> list[str]:
                cells = tr.findall("./w:tc", ns)
                vals: list[str] = []
                for tc in cells:
                    ts = [t.text for t in tc.findall(".//w:t", ns) if t.text]
                    vals.append("".join(ts).strip())
                return vals

            header = _row_cells(rows[0])
            if not header:
                continue
            md_table.append("| " + " | ".join(header) + " |")
            md_table.append("|" + "---|" * len(header))
            for tr in rows[1:]:
                md_table.append("| " + " | ".join(_row_cells(tr)) + " |")

            table_content = "\n".join(md_table).strip()
            if table_content:
                content_chunks.append({"type": "table", "content": table_content})

    # 5) 按文档顺序返回块列表
    return content_chunks


def parse_doc(file_path: str | Path) -> list[dict[str, str]]:
    """
    解析 **.doc**（Word 97–2003 二进制），输出与 ``parse_docx`` **相同的字典键**（``type`` / ``content``），
    便于与 ``chunks_to_markdown``、``parse_word_document`` 共用下游逻辑。

    **与 ``parse_docx`` 的差异（能否当同一种数据用）**

    - **结构层面**：可通用——均为 ``list[{"type": str, "content": str}]``，遍历方式一致。
    - **内容层面**：**不等价**。本函数依赖外部工具把整份 ``.doc`` 转成 **纯文本**：
      通常只返回 **一个** ``{"type": "text", "content": 全文}`` 块；**不会出现** ``type="table"``，
      段落边界、表格结构、与原文顺序的精细对应均弱于 OOXML 直解析。
    - 若要做与 ``.docx`` 同粒度的 RAG 分块，需在得到该大段 ``text`` 后 **再按段落/长度二次切分**。

    **格式样例**（常见情况为单块）::

        [
            {
                "type": "text",
                "content": "第一章 总则\\n\\n第一条 本规定适用于……\\n\\n（表格多为纯文本对齐，非 Markdown 表）",
            }
        ]

    **函数内各段代码作用**

    1. **路径与后缀**：仅接受 ``.doc``，与 ``.docx`` 分流，避免误调。
    2. **textutil（macOS）**：``textutil -convert txt -stdout`` 将二进制 ``.doc`` 转为 UTF-8 文本；
       成功且非空则 **立即返回** 单元素列表。
    3. **antiword（跨平台备选）**：若上一步失败（非 macOS、未装、或转换报错），再尝试命令行 ``antiword``。
    4. **失败聚合**：两步皆不可用或输出为空时抛出 ``RuntimeError``，提示检查文件与依赖。

    **部署提示**：Linux 服务器一般 **没有** textutil，请安装 **antiword** 或改用模块顶部文档中的其它 ``.doc`` 方案。

    :param file_path: ``.doc`` 文件路径。
    :raises ValueError: 后缀不是 ``.doc``。
    :raises RuntimeError: textutil 与 antiword 均无法得到有效文本。
    """
    # 1) 路径与后缀校验
    p = Path(file_path)
    if p.suffix.lower() != ".doc":
        raise ValueError(f"仅支持解析 .doc 文件，当前为: {p.name}")

    # 2) macOS：系统自带的 textutil 转纯文本
    try:
        resp = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(p)],
            capture_output=True,
            text=True,
            check=True,
        )
        text_content = (resp.stdout or "").strip()
        if text_content:
            return [{"type": "text", "content": text_content}]
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass

    # 3) 备选：antiword
    try:
        resp = subprocess.run(
            ["antiword", str(p)],
            capture_output=True,
            text=True,
            check=True,
        )
        text_content = (resp.stdout or "").strip()
        if text_content:
            return [{"type": "text", "content": text_content}]
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass

    # 4) 无可用后端
    raise RuntimeError(
        "DOC 解析失败：请确认文件未损坏；在 macOS 可用 textutil，或安装 antiword 后重试。"
    )


def parse_word_document(file_path: str | Path) -> list[dict[str, str]]:
    """
    统一入口：按后缀调用 ``parse_docx`` / ``parse_doc``，返回类型均为 ``list[dict[str, str]]``。

    注意：``.doc`` 路径多为单块纯文本；``.docx`` 可多块且含 ``table``，调用方勿假设块数量与类型分布一致。
    """
    p = Path(file_path)
    ext = p.suffix.lower()
    if ext == ".docx":
        return parse_docx(p)
    if ext == ".doc":
        return parse_doc(p)
    raise ValueError(f"仅支持 .doc/.docx，当前文件: {p.name}")


def _file_md5_hex(file_path: Path) -> str:
    with file_path.open("rb") as f:
        return hashlib.md5(f.read()).hexdigest()


# 与 ``chunks_to_markdown`` 中表格标题行一致，用于从已导出 MD 还原块类型
_MD_TABLE_HEADING_RE = re.compile(r"^###\s*表格块\s*\d+\s*$")


def markdown_from_word_export_to_parsed_blocks(md_text: str) -> list[dict[str, str]]:
    """
    将 ``export_doc_and_docx_to_markdown`` / ``chunks_to_markdown`` 生成的 Markdown **近似还原**
    为 ``parse_word_document`` 同形的块列表，便于再走 ``split_parsed_word_blocks``。

    规则：按连续空行（``\\n\\n+``）切段；若段首行为 ``### 表格块 {n}``，则该段为 ``type=table``，
    其余行为表格 Markdown；否则整段为 ``type=text``。

    说明：连续多个原文本块导出后仅表现为「多段用空行分隔」，还原后与「单段内含换行」在边界上可能无法区分，
    对 RAG 切片影响通常可接受。
    """
    md = (md_text or "").strip()
    if not md:
        return []
    segments = re.split(r"\n\s*\n+", md)
    blocks: list[dict[str, str]] = []
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        lines = seg.split("\n", 1)
        first = lines[0].strip()
        if _MD_TABLE_HEADING_RE.match(first):
            body = lines[1].strip() if len(lines) > 1 else ""
            if body:
                blocks.append({"type": "table", "content": body})
        else:
            blocks.append({"type": "text", "content": seg})
    return blocks


def split_parsed_word_blocks(
    parsed_blocks: list[dict[str, str]],
    splitter: RecursiveCharacterTextSplitter,
) -> list[tuple[str, int, str]]:
    """
    将 ``parse_word_document`` / ``parse_docx`` / ``parse_doc`` 返回的块列表再按 splitter 细切。

    返回 ``(切片文本, 原块序号从 1 起, 原块 type：text|table)``，顺序与原文块一致。
    """
    out: list[tuple[str, int, str]] = []
    block_idx = 0
    for raw in parsed_blocks:
        block_idx += 1
        ctype = raw.get("type") or "text"
        content = (raw.get("content") or "").strip()
        if not content:
            continue
        for piece in splitter.split_text(content):
            out.append((piece, block_idx, str(ctype)))
    return out


def build_chunks_for_word_document(
    file_path: str | Path,
    *,
    doc_idx: int = 1,
    file_hash: str | None = None,
    splitter: RecursiveCharacterTextSplitter | None = None,
    chunk_size: int = 500,
    chunk_overlap: int = 80,
    source_name: str | None = None,
    metadata_department: str = DEFAULT_WORD_CHUNK_METADATA_DEPARTMENT,
    metadata_update_time: str = DEFAULT_WORD_CHUNK_METADATA_UPDATE_TIME,
) -> list[dict[str, Any]]:
    """
    解析 ``.doc`` / ``.docx`` 并切成与 ``pdf_file_utils.build_chunks_for_pdf`` 同形的 RAG 块列表
    （``content`` + ``metadata``：``source_file``、``chunk_id``、``file_hash``、``department``、``update_time``）。

    Word 无稳定页码，额外写入 ``block_index``（对应解析块序号）、``block_type``（``text`` / ``table``）。
    ``file_hash`` 为 ``None`` 时对文件计算 MD5。
    """
    p = Path(file_path)
    parsed = parse_word_document(p)
    src = source_name if source_name is not None else p.name
    h = file_hash if file_hash is not None else _file_md5_hex(p)
    sp = splitter or RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    rows: list[dict[str, Any]] = []
    chunk_in_doc = 0
    for piece, block_idx, btype in split_parsed_word_blocks(parsed, sp):
        chunk_in_doc += 1
        rows.append(
            {
                "content": piece,
                "metadata": {
                    "source_file": src,
                    "chunk_id": f"doc_{doc_idx:02d}_ch_{chunk_in_doc}",
                    "file_hash": h,
                    "department": metadata_department,
                    "update_time": metadata_update_time,
                    "block_index": block_idx,
                    "block_type": btype,
                },
            }
        )
    return rows


def build_chunks_from_parsed_markdown_file(
    md_path: str | Path,
    *,
    doc_idx: int = 1,
    source_file: str | None = None,
    file_hash: str | None = None,
    original_doc_path: str | Path | None = None,
    splitter: RecursiveCharacterTextSplitter | None = None,
    chunk_size: int = 500,
    chunk_overlap: int = 80,
    metadata_department: str = DEFAULT_WORD_CHUNK_METADATA_DEPARTMENT,
    metadata_update_time: str = DEFAULT_WORD_CHUNK_METADATA_UPDATE_TIME,
) -> list[dict[str, Any]]:
    """
    读取由 ``chunks_to_markdown`` 写出的 ``.md``，还原块后按 ``split_parsed_word_blocks`` 切成 RAG 条，
    结构与 ``build_chunks_for_word_document`` 一致；``metadata`` 增加 ``content_basis: parsed_markdown``。

    ``file_hash``：默认对 **原始 Word**（``original_doc_path``）做 MD5；若未提供路径则对 ``.md`` 文件做 MD5。
    ``source_file``：逻辑来源文件名，默认使用 ``md_path`` 的文件名。
    """
    mp = Path(md_path)
    text = mp.read_text(encoding="utf-8")
    parsed = markdown_from_word_export_to_parsed_blocks(text)
    src = source_file if source_file is not None else mp.name
    if file_hash is not None:
        h = file_hash
    elif original_doc_path is not None and Path(original_doc_path).is_file():
        h = _file_md5_hex(Path(original_doc_path))
    else:
        h = _file_md5_hex(mp)

    sp = splitter or RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    rows: list[dict[str, Any]] = []
    chunk_in_doc = 0
    for piece, block_idx, btype in split_parsed_word_blocks(parsed, sp):
        chunk_in_doc += 1
        rows.append(
            {
                "content": piece,
                "metadata": {
                    "source_file": src,
                    "chunk_id": f"doc_{doc_idx:02d}_ch_{chunk_in_doc}",
                    "file_hash": h,
                    "department": metadata_department,
                    "update_time": metadata_update_time,
                    "block_index": block_idx,
                    "block_type": btype,
                    "content_basis": "parsed_markdown",
                    "parsed_markdown_file": mp.name,
                },
            }
        )
    return rows


def export_parsed_markdown_chunks_for_classified_docs(
    classified: dict[str, list[str]],
    base_dir: str | Path,
    parsed_md_dir: str | Path,
    chunks_json_dir: str | Path,
    *,
    chunk_size: int = 500,
    chunk_overlap: int = 80,
    metadata_department: str = DEFAULT_WORD_CHUNK_METADATA_DEPARTMENT,
    metadata_update_time: str = DEFAULT_WORD_CHUNK_METADATA_UPDATE_TIME,
) -> tuple[list[str], list[str]]:
    """
    对 ``classified["doc"]`` 中每个相对路径，读取 ``parsed_md_dir / {stem}.md``，切片后写入
    ``chunks_json_dir / {stem}_chunks.json``。

    返回 ``(成功写入的 json 路径列表, 失败或跳过原因列表)``。
    """
    base = Path(base_dir)
    pmd = Path(parsed_md_dir)
    out_dir = Path(chunks_json_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    ok_paths: list[str] = []
    errors: list[str] = []
    for doc_idx, rel in enumerate(classified.get("doc", []), start=1):
        stem = Path(rel).stem
        md_file = pmd / f"{stem}.md"
        if not md_file.is_file():
            errors.append(f"无 Markdown，跳过: {md_file}（rel={rel}）")
            continue
        orig = base / rel
        try:
            chunks = build_chunks_from_parsed_markdown_file(
                md_file,
                doc_idx=doc_idx,
                source_file=Path(rel).name,
                original_doc_path=orig if orig.is_file() else None,
                splitter=splitter,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                metadata_department=metadata_department,
                metadata_update_time=metadata_update_time,
            )
            out_json = out_dir / f"{stem}_chunks.json"
            out_json.write_text(
                json.dumps(chunks, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            ok_paths.append(str(out_json))
            print(
                f"[切片 JSON] {rel} -> {out_json.name}（{len(chunks)} 条）"
            )
        except Exception as e:
            errors.append(f"{rel} -> {e}")
            print(f"[切片失败] {rel} -> {e}")
    print(f"\n切片 JSON 输出目录：{out_dir}，成功 {len(ok_paths)}，异常/跳过 {len(errors)}")
    return ok_paths, errors


def chunks_to_markdown(chunks: list[dict[str, str]]) -> str:
    """将解析出的文本/表格块拼接为 Markdown。"""
    lines: list[str] = []
    for idx, ch in enumerate(chunks, start=1):
        ctype = ch.get("type", "text")
        content = (ch.get("content") or "").strip()
        if not content:
            continue
        if ctype == "table":
            lines.append(f"### 表格块 {idx}")
            lines.append(content)
        else:
            lines.append(content)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def export_doc_and_docx_to_markdown(
    classified: dict[str, list[str]],
    base_dir: str | Path,
    out_dir: str | Path,
) -> list[str]:
    """解析 doc/docx 并导出为 Markdown 文件。"""
    base = Path(base_dir)
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)

    ok_count = 0
    fail_count = 0
    failed_paths: list[str] = []
    for rel in classified.get("doc", []):
        src = base / rel
        try:
            chunks = parse_word_document(src)
            md_text = chunks_to_markdown(chunks)
            out_name = f"{Path(rel).stem}.md"
            out_file = out / out_name
            out_file.write_text(md_text, encoding="utf-8")
            ok_count += 1
        except Exception as e:
            fail_count += 1
            failed_paths.append(rel)
            print(f"[解析失败] {rel} -> {e}")
    print(f"\nWord 解析完成：成功 {ok_count}，失败 {fail_count}")
    print(f"Markdown 输出目录：{out}")
    return failed_paths
